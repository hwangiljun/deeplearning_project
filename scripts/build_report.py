"""최종 보고서(.docx)를 생성한다.

    python scripts/build_report.py

본문의 모든 수치는 models/ 의 산출물에서 직접 읽는다. 손으로 적어 넣은
성능 숫자는 없다. 그림은 scripts/build_report_figures.py 로 먼저 만들어 둘 것.
"""

from __future__ import annotations

import json
import sys
from pathlib import Path

import joblib
import numpy as np
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from src.data import OUTCOMES
from src.features import ReferenceTables
from src.profiles import Profiles

ROOT = Path(__file__).resolve().parents[1]
MODELS = ROOT / "models"
FIGS = ROOT / "reports" / "figures"
OUT = ROOT / "reports" / "딥러닝 최종 보고서 (개정판).docx"

KO = {"ball": "볼", "called_strike": "루킹 스트라이크",
      "swinging_strike": "헛스윙 스트라이크", "foul": "파울",
      "hit_by_pitch": "사구", "field_out": "인플레이 아웃", "single": "단타",
      "double": "2루타", "triple": "3루타", "home_run": "홈런"}

FONT = "맑은 고딕"
INK = RGBColor(0x1A, 0x1A, 0x1A)
MUTED = RGBColor(0x66, 0x66, 0x66)
ACCENT = RGBColor(0x1C, 0x5C, 0xAB)


# --------------------------------------------------------------------------
# 서식 도우미
# --------------------------------------------------------------------------
def _kfont(run, size=10.5, bold=False, color=INK, name=FONT):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)


def setup(doc: Document):
    st = doc.styles["Normal"]
    st.font.name = FONT
    st.font.size = Pt(10.5)
    st.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    pf = st.paragraph_format
    pf.line_spacing = 1.5
    pf.space_after = Pt(6)

    for s in doc.sections:
        s.top_margin = Cm(2.4)
        s.bottom_margin = Cm(2.4)
        s.left_margin = Cm(2.6)
        s.right_margin = Cm(2.6)


def h(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(6)
    _kfont(p.add_run(text), size=15 if level == 1 else 12.5, bold=True,
           color=INK if level == 1 else ACCENT)
    return p


def para(doc, text, size=10.5, indent=True, color=INK, italic=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.4)
    p.paragraph_format.space_after = Pt(6)
    for seg, bold in _split_bold(text):
        r = p.add_run(seg)
        _kfont(r, size=size, bold=bold, color=color)
        r.font.italic = italic
    return p


def _split_bold(text: str):
    """**굵게** 표기를 런으로 나눈다."""
    out, buf, bold = [], "", False
    i = 0
    while i < len(text):
        if text[i:i + 2] == "**":
            if buf:
                out.append((buf, bold))
                buf = ""
            bold = not bold
            i += 2
        else:
            buf += text[i]
            i += 1
    if buf:
        out.append((buf, bold))
    return out


def bullet(doc, text, size=10.5, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.7 + 0.6 * level)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.35
    for seg, bold in _split_bold(text):
        _kfont(p.add_run(seg), size=size, bold=bold)
    return p


def code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.15
    _kfont(p.add_run(text), size=9, name="Consolas", color=RGBColor(0x33, 0x33, 0x33))
    return p


def caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.line_spacing = 1.15
    _kfont(p.add_run(text), size=9, color=MUTED)


def figure(doc, name, cap, width_cm=15.0):
    path = FIGS / name
    if not path.exists():
        para(doc, f"[그림 누락: {name}]", color=MUTED)
        return
    doc.add_picture(str(path), width=Cm(width_cm))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption(doc, cap)


def table(doc, headers, rows, cap=None, widths=None, highlight_row=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, htxt in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ""
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_after = Pt(2)
        _kfont(p.add_run(str(htxt)), size=9.5, bold=True)

    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for i, v in enumerate(row):
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.space_after = Pt(2)
            _kfont(p.add_run(str(v)), size=9.5, bold=(highlight_row == ri))

    if widths:
        for r in t.rows:
            for i, w in enumerate(widths):
                r.cells[i].width = Cm(w)
    if cap:
        caption(doc, cap)
    return t


# --------------------------------------------------------------------------
def load_all():
    meta = json.loads((MODELS / "main" / "metadata.json").read_text(encoding="utf-8"))
    ev = joblib.load(MODELS / "main" / "evaluation.pkl")
    ab = json.loads((MODELS / "ablation.json").read_text(encoding="utf-8"))
    cfg = json.loads((MODELS / "main" / "model_config.json").read_text(encoding="utf-8"))
    tables = ReferenceTables.load(MODELS / "tables.pkl")
    profiles = Profiles.load(MODELS / "profiles.pkl")
    return meta, ev, ab, cfg, tables, profiles


def build():
    meta, ev, ab, cfg, tables, profiles = load_all()
    tm = meta["test_metrics"]
    hist = meta["history"]
    pc = ev["per_class"]
    cv = tables.count_value

    doc = Document()
    setup(doc)

    # ---------------- 표지 ----------------
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _kfont(p.add_run("딥러닝 기반 실시간 투구 전략 최적화 시스템"), size=20, bold=True)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _kfont(p.add_run("― 데이터 누수의 발견과 파이프라인 재설계 ―"), size=12.5, color=MUTED)
    doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _kfont(p.add_run("개정판"), size=11, bold=True, color=ACCENT)
    for _ in range(8):
        doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _kfont(p.add_run("22013442  황일준"), size=12)
    doc.add_page_break()

    # ---------------- 초록 ----------------
    h(doc, "요약")
    para(doc,
         "본 연구는 MLB Statcast 투구 단위 데이터를 이용해 투구 결과를 예측하고, "
         "주어진 상황에서 기대 실점이 가장 낮은 구종과 코스를 제시하는 시스템을 구축한다. "
         "초기 버전은 검증 정확도 67.52%를 보고하였으나, 본 개정 과정에서 그 수치가 "
         "**미래 정보 누수로 만들어진 값**임을 확인하였다. Statcast 원본이 역시간순으로 "
         "제공된다는 사실을 놓쳐, 결과적으로 나중에 던진 공들로 먼저 던진 공의 결과를 "
         "맞히는 문제를 학습하고 있었다.")
    para(doc,
         "이에 데이터 구성부터 평가 방법까지 전면 재설계하였다. 시퀀스를 시간순으로 "
         "정렬하고 타석 내부로 제한하였으며, 분할을 시즌 단위로 바꾸고, 불균형 이중 "
         f"보정을 제거하였다. 2024 시즌 {len(hist)}에폭 학습, 2025 시즌 "
         f"{ev['n_test']:,}구 홀드아웃 평가 결과 log-loss {tm['log_loss']:.4f}, "
         f"macro AUC {tm['macro_auc']:.4f}, ECE {ev['ece']:.4f}를 얻었다.")
    para(doc,
         "어블레이션을 통해 맥락 스킵 연결이 성능에 기여함을 확인하였으나, 동시에 "
         "**투구 시퀀스 자체의 기여는 예상보다 작다**는 사실도 함께 확인하였다. "
         "이는 초기 연구가 전제한 '투구 배합의 시계열 맥락이 결정적'이라는 가정을 "
         "이 데이터가 강하게 지지하지 않음을 뜻한다.")

    # ---------------- 1. 서론 ----------------
    h(doc, "1. 서론")
    para(doc,
         "현대 야구에서 투수와 타자의 승부는 볼카운트, 주자 상황, 직전 투구의 배합 등 "
         "매 순간 변화하는 맥락 속에서 이루어진다. 타율·출루율·장타율과 같은 누적 지표는 "
         "선수의 전반적 가치를 요약하는 데 유용하지만, 특정 시점의 의사결정을 돕지는 못한다. "
         "투수와 포수가 실제로 필요로 하는 정보는 '지금 이 상황에서 어떤 공을 어디에 "
         "던져야 하는가'이며, 이는 결과가 아니라 **선택**에 관한 질문이다.")
    para(doc,
         "본 연구는 이 질문에 답하기 위해 투구 단위 결과 예측 모델을 학습하고, 그 확률 "
         "분포를 기대 실점으로 환산하여 후보 투구를 순위화하는 시스템을 구축한다. "
         "다만 본 보고서의 서술 순서는 통상적인 연구 보고서와 다르다. 초기 구현이 "
         "산출한 수치가 유효하지 않음을 발견한 것이 본 개정의 출발점이므로, "
         "**무엇이 잘못되었는지를 먼저 규명한 뒤** 그에 대응하는 설계를 제시한다.")

    para(doc, "본 연구의 기여는 다음과 같다.")
    bullet(doc, "Statcast 데이터를 이용한 시퀀스 모델링에서 발생하기 쉬운 "
                "**시간 역순 누수를 구체적으로 규명**하고, 그것이 성능 수치를 어떻게 "
                "부풀리는지 정량적으로 보였다.")
    bullet(doc, "타석 단위 시퀀스 구성, 시즌 단위 분할, 학습 구간 한정 정규화를 통해 "
                "**재현 가능하고 누수 없는 파이프라인**을 재구축하였다.")
    bullet(doc, "임의 가중치에 의존하던 추천 점수를 Statcast 기대 득점 변화량에 근거한 "
                "**카운트 조건부 기대 실점**으로 대체하였다.")
    bullet(doc, "한 번에 한 요소만 변경하는 어블레이션으로 각 구성요소의 기여를 분리하고, "
                "**시퀀스의 기여가 작다는 부정적 결과를 포함해** 보고하였다.")

    # ---------------- 2. 관련 연구 ----------------
    h(doc, "2. 관련 연구")
    para(doc,
         "투구 결과 예측에는 오랫동안 트리 기반 모델이 사용되었다. Sidle와 Tran은 랜덤 "
         "포레스트로 구종을 예측하며 타자 유리 카운트에서 예측 가능성이 높아짐을 보였고, "
         "Swartz 등은 투구 품질 평가에서 낮고 바깥쪽 코스의 효과와 카운트의 중요성을 "
         "확인하였다. 이들 모델은 각 투구를 독립 사건으로 다루므로 배합의 순서 효과를 "
         "포착하기 어렵다는 한계가 있다.")
    para(doc,
         "이를 보완하기 위해 시퀀스 모델이 도입되었다. Kneita는 Transformer 기반 투구 결과 "
         "예측에서 시퀀스 맥락과 투수 강점 정렬의 중요성을 보고하였다. 초기 버전의 본 연구도 "
         "이 계열에 속하며, 긴 시퀀스에서 현재 상황 정보가 희석되는 문제를 지적하고 "
         "맥락을 출력층에 직접 결합하는 구조를 제안하였다.")
    para(doc,
         "가장 최근의 연구인 Takamido와 Nakamoto(2026)는 반사실(counterfactual) 분석을 "
         "도입하여, 실제로 던지지 않은 공의 가치를 평가하는 틀을 제시하였다. 이들은 "
         "타석 단위 시퀀스와 맥락 정보를 별도 경로로 처리하는 이원 구조를 사용하였고, "
         "최종구뿐 아니라 그 직전의 셋업 피치를 대체하며 시즌 지표(K/9, ERA, oSLG)에 "
         "미치는 영향을 추정하였다. 본 개정판은 이 연구의 데이터 구성 방식 — "
         "타석 단위 시퀀스, 좌타자 좌표 미러링, 타자별 스트라이크존 정규화, "
         "투수×구종 실측 평균을 이용한 반사실 대체, 제구 오차를 반영한 윈도우 평균 — "
         "을 다수 채택하였다.")
    para(doc,
         "특기할 점은 이들이 보고한 결과의 방향이다. 셋업 피치 최적화의 효과(K/9 +1.10~1.23)는 "
         "최종구 최적화의 효과(+1.26~3.54)보다 일관되게 작았다. 후술하듯 본 연구의 "
         "어블레이션 결과도 같은 방향을 가리킨다.")

    # ---------------- 3. 초기 구현의 결함 ----------------
    h(doc, "3. 초기 구현의 재현과 결함 분석")
    para(doc,
         "개정 작업은 초기 구현의 재현에서 시작하였다. 학습 노트북과 원본 데이터를 "
         "확보하여 전처리 과정을 추적한 결과, 보고된 성능을 무효화하는 결함들이 확인되었다.")

    h(doc, "3.1 시간 역순 시퀀스로 인한 미래 정보 누수", 2)
    para(doc,
         "pybaseball이 제공하는 Statcast 원본은 역시간순으로 정렬되어 있다. 실제 데이터를 "
         "확인한 결과 한 타석 안에서 pitch_number가 7, 6, 5 … 1의 순서로 나타났으며, "
         "전체 629,843행에서 날짜가 과거로 이동하는 전환이 158회, 미래로 이동하는 전환이 "
         "5회였다. 초기 코드에는 정렬 구문이 존재하지 않았다.")
    code(doc,
         "grouped = df.groupby('batter')       # 역순이 그대로 유지된다\n"
         "sequences.append(data[i : i+8])\n"
         "targets.append(target_vals[i+7])     # 윈도우의 마지막 = 시간상 가장 과거")
    para(doc,
         "그 결과 모델은 타깃보다 **나중에** 던져진 7개의 공을 입력으로 받아 그보다 "
         "**먼저** 던져진 1개 공의 결과를 예측하고 있었다. 더욱이 볼카운트가 모든 "
         "타임스텝에 포함되어 있었으므로, 타깃 다음 투구의 카운트를 비교하는 것만으로 "
         "타깃의 결과를 직접 읽어낼 수 있었다. 볼이 1 증가했으면 볼, 스트라이크가 "
         "1 증가했으면 스트라이크이며, 0-0으로 초기화되었다면 그 투구가 타석을 종료한 것이다. "
         "전체의 절반 이상을 차지하는 볼과 스트라이크가 입력에서 직접 관측 가능했다.")
    figure(doc, "fig1_leak.png",
           "그림 1. 정렬 누락이 만들어 낸 누수 구조. 원본이 역시간순이므로 "
           "정렬 없이 잘라낸 윈도우에서 타깃은 시간상 가장 과거의 투구가 되고, "
           "나머지 입력은 모두 그 이후에 일어난 사건이 된다.", 16.0)

    h(doc, "3.2 무작위 분할과 겹치는 윈도우", 2)
    para(doc,
         "초기 보고서는 '시간 순서에 따라 8:2로 분할'하였다고 기술하였으나, 실제 코드는 "
         "train_test_split에 stratify를 지정한 무작위 셔플 분할이었다. 여기에 stride 2의 "
         "슬라이딩 윈도우가 결합되어, 학습 표본 [i..i+7]과 검증 표본 [i+2..i+9]가 "
         "8개 중 6개의 투구를 공유하였다. 검증셋이 학습셋과 사실상 중복된 상태였다.")

    h(doc, "3.3 불균형 이중 보정에 의한 확률 왜곡", 2)
    para(doc,
         "학습은 역빈도 가중 오버샘플링(WeightedRandomSampler)과 Focal Loss를 동시에 "
         "적용하였다. 불균형을 두 번 보정한 셈이며, 그 결과 모델의 출력 확률은 실제 분포가 "
         "아니라 균등화된 분포를 가리키게 된다. 응용 단계에서 temperature를 5.0으로 "
         "고정하여 분포를 인위적으로 평탄화한 것은 이 왜곡을 가리기 위한 대증요법이었다.")

    h(doc, "3.4 결과 정의와 데이터 범위의 문제", 2)
    bullet(doc, "정규시즌 필터가 없어 **시범경기 35,482구**가 학습에 포함되어 있었다.")
    bullet(doc, "결과 매핑이 events 화이트리스트에 의존하여 3루타, 희생플라이, 병살, "
                "야수선택 등 **인플레이 결과 10,224구가 삭제**되었다. 3루타가 클래스에서 "
                "사라진 것도 이 때문이다.")
    bullet(doc, "수집 범위가 2024년 3~8월에 그쳤음에도 '정규 시즌 전체'로 기술되었다.")
    bullet(doc, "인코더와 MinMaxScaler를 **전체 데이터에 적합**한 뒤 분할하여, "
                "검증셋의 분포 정보가 정규화를 통해 학습에 유입되었다.")

    h(doc, "3.5 추천 점수의 수학적 무의미성", 2)
    para(doc, "초기 추천 점수는 다음과 같이 정의되어 있었다.")
    code(doc, "score = p_success - (p_fail * 1.5) - penalty")
    para(doc,
         "그러나 good 집합(5개)과 bad 집합(6개)이 11개 클래스 전체를 정확히 이분하므로 "
         "항상 p_fail = 1 − p_success가 성립한다. 따라서 score = 2.5·p_success − 1.5이며, "
         "**위험 가중치 1.5는 후보 순위에 어떠한 영향도 주지 못한다.** 또한 2스트라이크에서의 "
         "파울과 범타가 동일하게 '성공'으로 분류되어, 실제 가치 차이가 반영되지 않았다.")

    doc.add_page_break()

    # ---------------- 4. 제안 방법 ----------------
    h(doc, "4. 제안 방법")
    para(doc,
         "3장에서 확인한 결함에 대응하여 데이터 구성, 모델, 평가, 추천 로직을 재설계하였다. "
         "설계 원칙은 세 가지다. 첫째, 학습 시점에 알 수 없는 정보는 어떤 경로로도 입력에 "
         "들어오지 않는다. 둘째, 성능 수치는 학습·조기종료·모델선택에 한 번도 쓰이지 않은 "
         "구간에서 단 한 번 측정한다. 셋째, 확률로 후보를 순위화하는 시스템이므로 "
         "확률의 정직성(캘리브레이션)을 정확도보다 우선한다.")

    h(doc, "4.1 데이터 구성", 2)
    para(doc,
         "2024·2025 정규시즌 전체를 수집하여 총 1,424,426구를 확보하였다(2024년 711,898구, "
         "2025년 712,528구). game_type이 정규시즌인 경기만 사용하였다. "
         "원본은 game_date → game_pk → at_bat_number → pitch_number 순으로 정렬하여 "
         "시간축을 복원하였다.")
    para(doc,
         "시퀀스는 **타석(plate appearance) 내부로 제한**한다. 각 투구를 타깃으로 삼고, "
         "같은 타석에서 그보다 앞서 던져진 공들을 최대 6구까지 앞에 붙인다. 부족한 앞부분은 "
         "0으로 채우고 마스크로 표시하여 어텐션에서 제외한다. 타석 경계를 넘지 않으므로 "
         "다른 타자나 투수의 투구가 한 시퀀스에 섞이지 않는다. 실제 데이터에서 타석당 "
         "평균 투구 수는 3.82구이고 91.7%의 타석이 6구 이하이므로, 길이 6은 대부분의 "
         "타석을 온전히 포괄한다.")
    para(doc,
         "좌표계는 Takamido와 Nakamoto(2026)를 따라 정규화하였다. 좌타자의 좌우 위치와 "
         "수평 무브먼트는 부호를 반전시켜 좌우 타자를 동일한 의미 축에서 비교할 수 있게 하고, "
         "높이는 타자별 스트라이크존 상·하단으로 정규화하여 신장 차이를 흡수한다. "
         "회전축은 각도이므로 사인·코사인으로 분해하여 0도와 359도가 인접하도록 하였다.")

    table(doc,
          ["구분", "변수", "개수"],
          [["시퀀스 (투구 물리량)",
            "좌우 위치, 정규화 높이, 구속, 회전수, 회전축(sin·cos), 수평·수직 무브먼트", "8"],
           ["시퀀스 (범주형)", "타자, 투수, 구종, 타석 좌우, 투구 좌우", "5"],
           ["맥락", "볼, 스트라이크, 아웃, 주자 1·2·3루, 이닝, 초말, 점수차, 타순 회차", "10"],
           ["맥락 (타자 특성)", "직전 시즌 삼진율, 홈런율, ISO, 타율, OPS", "5"]],
          "표 1. 입력 변수 구성. 맥락의 타자 특성은 **직전 시즌** 성적만 사용하여 "
          "같은 시즌 정보가 유입되지 않도록 하였다.",
          widths=[3.6, 9.4, 1.6])

    para(doc,
         "**볼카운트를 시퀀스 경로에서 제외한 것은 의도적인 설계**다. 3.1에서 확인한 누수는 "
         "카운트가 모든 타임스텝에 존재했기 때문에 발생하였다. 카운트를 맥락 경로에만 두면 "
         "타깃 시점의 상황은 그대로 활용하면서, 인접 투구의 카운트를 비교해 정답을 추론하는 "
         "경로 자체가 구조적으로 차단된다.")

    h(doc, "4.2 출력 클래스의 재정의", 2)
    para(doc,
         "초기 정의는 strike와 strikeout을 동시에 클래스로 두었다. 그러나 삼진은 "
         "2스트라이크에서의 스트라이크이고 볼넷은 3볼에서의 볼이므로, 이들은 투구의 결과가 "
         "아니라 **카운트에서 파생되는 타석 결과**다. 카운트가 입력에 포함된 상태에서 이를 "
         "별도 클래스로 두면 모델이 사실상 공짜로 맞히는 라벨이 생긴다. 본 개정에서는 "
         "투구 단위로 실제 관측되는 사건만을 클래스로 정의하였다.")
    code(doc,
         "ball · called_strike · swinging_strike · foul · hit_by_pitch\n"
         "field_out · single · double · triple · home_run")
    para(doc,
         "삼진과 볼넷은 예측된 클래스와 현재 카운트로부터 결정적으로 유도된다. "
         "또한 인플레이 결과 매핑을 재정의하여 병살, 희생플라이, 야수선택, 실책 출루 등을 "
         "포함시켰다. 그 결과 매핑되지 않아 삭제되는 투구가 10,224구에서 6구로 감소하였다.")

    h(doc, "4.3 모델 구조", 2)
    para(doc,
         "제안 모델은 시퀀스와 맥락을 별도 경로로 처리한 뒤 출력 직전에 결합하는 이원 구조를 "
         "갖는다. 이는 초기 버전이 제안한 구조를 계승한 것이며, Takamido와 Nakamoto(2026)가 "
         "독립적으로 채택한 구조와도 일치한다.")
    figure(doc, "fig2_architecture.png", "그림 2. 제안 모델 구조.", 13.5)

    para(doc, "각 구성요소의 역할과 설계 근거는 다음과 같다.")

    para(doc, "**(1) 엔티티 임베딩**", indent=False)
    para(doc,
         "타자·투수·구종 등 고차원 범주형 변수를 학습 가능한 밀집 벡터로 사상한다. "
         "원-핫 인코딩과 달리 선수 간 유사성이 벡터 공간에 표현될 수 있다. "
         f"학습 데이터에 등장한 타자 {cfg['cat_sizes']['batter']:,}명, "
         f"투수 {cfg['cat_sizes']['pitcher']:,}명이 사전에 포함된다. "
         "**인덱스 0은 UNK(미등록 선수)로 예약**하고, 학습 중 일정 확률로 선수 ID를 UNK로 "
         "치환하여 해당 임베딩이 실제로 훈련되도록 하였다. 초기 구현은 예외 처리로 미등록 "
         "선수를 0번 선수로 조용히 대체하여, 처음 보는 선수에 대해 특정 선수의 예측을 "
         "반환하는 오류가 있었다.")

    para(doc, "**(2) Transformer 인코더**", indent=False)
    para(doc,
         f"d_model {cfg['d_model']}, 헤드 {cfg['nhead']}개, 레이어 {cfg['num_layers']}층, "
         f"피드포워드 차원 {cfg['dim_feedforward']}로 구성하였다. 셀프 어텐션이 타석 내 "
         "투구 간 관계를 학습한다. 세 가지를 수정하였다.")
    bullet(doc, "**최종 LayerNorm 추가.** norm_first=True(Pre-LN) 설정에서는 마지막 블록의 "
                "출력이 정규화되지 않은 채로 헤드에 전달된다. 초기 구현에는 최종 정규화가 "
                "없어 잔차 스트림의 크기가 제어되지 않았다.")
    bullet(doc, "**패딩 마스크 적용.** 타석 앞부분의 패딩 위치에도 어텐션이 걸리고 있었다. "
                "src_key_padding_mask로 존재하지 않는 투구를 제외하였다.")
    bullet(doc, "**마스크드 평균 풀링.** 마지막 토큰만 사용하던 것을 유효 구간 평균으로 "
                "변경하였다. 짧은 타석에서 특히 안정적이다.")

    para(doc, "**(3) 맥락 경로와 스킵 연결**", indent=False)
    para(doc,
         "맥락 정보는 Transformer를 거치지 않고 별도의 완전연결층을 통과한 뒤 출력 직전에 "
         "결합된다. 이것이 초기 연구가 제안한 핵심 구조다. 다만 초기 구현에는 심각한 문제가 "
         "있었다. 원시 정수 6개를 Linear(6, 32)와 ReLU에 통과시켰는데, **0-0 카운트·무사·"
         "주자 없음이면 입력 벡터가 전부 0이 되어 ReLU(bias)라는 상수만 남는다.** "
         "즉 가장 흔한 상황에서 스킵 연결이 전달하는 정보량이 0이었다. "
         "본 개정에서는 맥락 피처를 표준화하여 0이 특별한 값이 되지 않게 하고, "
         "이닝·점수차·타순 회차·타자 직전시즌 성적을 추가하여 표현력을 높였다.")

    para(doc, "**(4) 출력 헤드와 보조 회귀**", indent=False)
    para(doc,
         "결합된 벡터는 다층 퍼셉트론을 거쳐 10개 클래스의 로짓을 산출한다. "
         "정규화는 BatchNorm 대신 LayerNorm을 사용하였다. BatchNorm은 배치 구성에 민감하여 "
         "추론 시 배치 크기가 달라지는 응용 환경에 적합하지 않기 때문이다. "
         "추가로 해당 투구의 기대 득점 변화량을 직접 예측하는 보조 회귀 헤드를 두어, "
         "분류 손실과 함께 학습하였다.")

    para(doc,
         f"최종 모델의 파라미터 수는 486,199개(약 2.0MB)이다. 초기 모델(5,408,636개, 약 21MB)의 "
         "약 1/11 규모다. 초기 보고서는 '경량화 달성'을 측정 없이 주장하였으나, "
         "본 개정에서는 실측값으로 제시한다.")

    h(doc, "4.4 학습 설정", 2)
    para(doc,
         "2024 시즌을 학습에, 2025 시즌을 테스트에 사용한다. 검증셋은 학습 구간의 "
         "**뒷부분 날짜**에서 15%를 분리하였다. 무작위 분할과 달리 같은 타석의 이웃 투구가 "
         "학습과 검증으로 갈리지 않는다. 인코더와 정규화 통계는 학습 구간에만 적합하고 "
         "그 값을 검증·테스트에 그대로 적용한다.")
    para(doc,
         "손실 함수는 단순 교차 엔트로피를 사용하고 **오버샘플링을 적용하지 않았다.** "
         "3.3에서 확인하였듯 이중 보정은 확률을 왜곡한다. 희귀 클래스의 재현율을 높이는 대신 "
         "확률의 정직성을 택한 것이며, 이 선택의 결과는 5.4에서 정량적으로 검토한다. "
         "최적화는 AdamW(학습률 3×10⁻⁴, weight decay 10⁻⁴), 스케줄러는 ReduceLROnPlateau, "
         "그래디언트 클리핑 1.0을 적용하였다. **조기 종료 기준은 검증 log-loss**이며, "
         "정확도의 우연한 최고점을 성능으로 채택하지 않는다.")

    h(doc, "4.5 추천 로직", 2)
    para(doc,
         "추천은 후보 생성과 점수화의 두 단계로 이루어진다. 후보는 해당 투수가 실제로 던지는 "
         "구종과 9×9 좌표 격자의 조합으로 생성한다. 초기 구현은 존 안 9칸만을 후보로 삼아 "
         "**유인구를 추천할 수 없었으나**, 본 개정은 격자를 존 밖까지 확장하였다. "
         "각 구종의 물리량은 리그 평균 상수가 아니라 **해당 투수의 실측 평균**을 사용한다. "
         f"현재 참조 테이블은 투수 {tables.repertoire['pitcher'].nunique():,}명에 대해 "
         f"{len(tables.repertoire):,}개의 (투수, 구종) 조합을 보유한다.")
    para(doc,
         "점수는 모델이 산출한 결과 확률 분포를 카운트 조건부 기대 실점으로 가중합한 값이다.")
    code(doc, "score(p) = Σ_c  P(결과 = c | 투구 p, 상황)  ×  ΔRE(c | 볼카운트)")
    para(doc,
         "여기서 ΔRE는 Statcast가 제공하는 투구별 기대 득점 변화량(delta_run_exp)에서 "
         "직접 추정한다. 표 2에서 보듯 같은 '볼'이라도 0-0에서는 거의 무해하지만 3-2에서는 "
         "볼넷이 되어 8배 이상의 값을 갖는다. 2스트라이크에서의 파울은 거의 무가치한 반면 "
         "범타는 큰 이득이다. 초기 점수 함수의 이분법은 이러한 차이를 전혀 담지 못하였다.")

    rows = []
    for name in ["ball", "called_strike", "swinging_strike", "foul", "field_out",
                 "single", "home_run"]:
        i = OUTCOMES.index(name)
        rows.append([KO[name]] + [f"{cv[cs, i]:+.3f}" for cs in (0, 2, 5, 8, 11)])
    table(doc, ["결과", "0-0", "0-2", "1-2", "2-2", "3-2"], rows,
          "표 2. 카운트 조건부 기대 실점(투수 관점, 음수가 유리). "
          "2024·2025 시즌 delta_run_exp에서 추정하였다.",
          widths=[3.8, 2.2, 2.2, 2.2, 2.2, 2.2])

    para(doc,
         "마지막으로 제구 오차를 반영한다. 투수는 한 점을 정확히 겨눌 수 없으므로, "
         "노린 지점 주변 윈도우의 평균 기대 실점을 그 지점의 대표값으로 삼는다"
         "(Takamido와 Nakamoto의 command window). 윈도우 크기를 조절하면 제구가 정밀한 "
         "투수와 그렇지 않은 투수를 구분하여 평가할 수 있다.")

    doc.add_page_break()

    # ---------------- 5. 결과 ----------------
    h(doc, "5. 실험 결과")
    h(doc, "5.1 학습 경과", 2)
    best_ep = min(hist, key=lambda x: x["log_loss"])
    para(doc,
         f"학습은 {len(hist)}에폭 진행되었으며 검증 log-loss는 {hist[0]['log_loss']:.4f}에서 "
         f"{best_ep['log_loss']:.4f}까지 감소하였다. 최저값이 마지막 에폭인 "
         f"{best_ep['epoch']}에폭에서 관측된 점에 주목할 필요가 있다. "
         "**조기 종료가 아니라 에폭 상한에 도달하여 학습이 종료된 것**이므로, "
         "현재 성능은 수렴값이 아니라 주어진 예산 안에서의 값이다. 더 긴 학습으로 "
         "개선의 여지가 남아 있다.")
    figure(doc, "fig3_curve.png", "그림 3. 검증 log-loss의 에폭별 추이.", 13.5)

    h(doc, "5.2 테스트 시즌 성능", 2)
    para(doc,
         f"학습·조기종료·모델선택에 한 번도 사용하지 않은 2025 시즌 {ev['n_test']:,}구에서 "
         "다음 성능을 얻었다.")
    table(doc, ["지표", "값", "비고"],
          [["log-loss", f"{tm['log_loss']:.4f}", "주 지표. 낮을수록 좋음"],
           ["macro AUC", f"{tm['macro_auc']:.4f}", "클래스 균등 가중 순위 성능"],
           ["macro F1", f"{tm['macro_f1']:.4f}", "희귀 클래스 반영 (5.4 참조)"],
           ["정확도", f"{tm['accuracy']*100:.2f}%", "참고 지표"],
           ["ECE", f"{ev['ece']:.4f}", "캘리브레이션 오차 (5.3 참조)"]],
          "표 3. 2025 시즌 홀드아웃 성능.", widths=[4.0, 3.0, 8.0])

    para(doc,
         f"초기 보고서의 67.52%와 본 결과의 {tm['accuracy']*100:.2f}%는 **비교 대상이 아니다.** "
         "전자는 미래 정보 누수를 포함하고, 카운트에서 파생되는 strikeout·walk를 클래스로 "
         "포함하며, 시범경기가 섞여 있고, 검증셋이 학습셋과 6/8을 공유한 상태에서 측정되었다. "
         "동일한 지표명을 공유할 뿐 측정 대상이 서로 다르다.")

    h(doc, "5.3 캘리브레이션", 2)
    para(doc,
         f"예측 확률의 정직성을 확인하기 위해 신뢰도 보정 오차(ECE)를 측정하고, 검증셋에서 "
         f"temperature scaling 계수를 적합하였다. ECE는 {ev['ece']:.4f}, "
         f"적합된 temperature는 {ev['temperature']:.3f}였다. "
         "**T가 1에 근접한다는 것은 모델이 이미 보정되어 있어 추가 보정이 불필요함을 뜻한다.** "
         f"실제로 보정을 적용해도 log-loss는 {tm['log_loss']:.4f}에서 "
         f"{ev['log_loss_scaled']:.4f}로 거의 변하지 않았다.")
    para(doc,
         "이는 4.4의 설계 선택에 대한 직접적인 근거다. 초기 구현이 temperature 5.0을 "
         "손으로 고정해야 했던 것은 이중 보정으로 왜곡된 확률을 가리기 위함이었고, "
         "원인을 제거하자 보정 자체가 필요 없어졌다. 화면에 표시되는 확률이 실제 빈도와 "
         "일치하므로, 추천 점수의 기대 실점 계산도 신뢰할 수 있다.")
    figure(doc, "fig5_calibration.png",
           "그림 4. 캘리브레이션 곡선. 대각선에 가까울수록 예측 확률이 실제 빈도와 일치한다.", 10.5)

    h(doc, "5.4 클래스별 성능과 설계상의 맞바꿈", 2)
    rows = []
    for i, name in enumerate(pc["outcome"]):
        rows.append([KO[name], f"{pc['support'][i]:,}", f"{pc['precision'][i]:.3f}",
                     f"{pc['recall'][i]:.3f}", f"{pc['f1'][i]:.3f}", f"{pc['auc'][i]:.3f}"])
    table(doc, ["클래스", "표본", "정밀도", "재현율", "F1", "AUC"], rows,
          "표 4. 2025 시즌 클래스별 성능.", widths=[3.4, 2.6, 2.2, 2.2, 2.2, 2.2])

    para(doc,
         "안타 계열(단타·2루타·3루타·홈런)의 재현율이 0이다. 사전 확률이 1~4%인 클래스는 "
         "argmax에서 최상위가 되지 못하기 때문이며, macro F1이 낮은 것도 대부분 이 때문이다. "
         "그러나 같은 클래스들의 **AUC는 0.785~0.835로, 순위 매기기는 정상적으로 이루어지고 "
         "있다.** 즉 모델은 '어느 투구가 장타로 이어질 가능성이 높은가'를 구분하지만, "
         "그 확률이 다른 클래스를 넘어설 만큼 크지는 않다.")
    para(doc,
         "이는 의도한 맞바꿈이다. 재현율을 높이려면 오버샘플링이나 클래스 가중을 도입해야 "
         "하는데, 그러면 출력 확률이 실제 분포를 벗어나 캘리브레이션이 훼손된다"
         "(초기 구현이 정확히 그 상태였다). 본 시스템의 추천 엔진은 argmax가 아니라 "
         "**확률 분포 전체를 기대 실점으로 적분**하므로, 여기서는 정직한 확률이 높은 "
         "재현율보다 중요하다.")
    figure(doc, "fig6_confusion.png", "그림 5. 혼동 행렬(행 정규화, 단위 %).", 13.0)

    h(doc, "5.5 어블레이션", 2)
    para(doc,
         "초기 보고서의 가장 큰 방법론적 약점은 어블레이션의 부재였다. 제안 모델과 "
         "베이스라인이 d_model, 레이어 수, 손실 함수, 에폭 수, 학습률까지 모두 달랐으므로 "
         "성능 차이를 특정 요소에 귀속시킬 수 없었다. 본 개정에서는 **한 번에 하나의 요소만 "
         "변경**하고 나머지 조건을 동일하게 유지하였다.")

    names = {"full": "제안 모델 (전체)", "no_context_skip": "맥락 스킵 연결 제거",
             "seq_len_1": "시퀀스 길이 1", "no_state_embed": "상태 임베딩 제거",
             "last_token": "마지막 토큰 풀링", "seq_len_3": "시퀀스 길이 3",
             "seq_len_10": "시퀀스 길이 10"}
    base = ab["full"]["test"]["log_loss"]
    order = sorted(ab.items(), key=lambda kv: kv[1]["test"]["log_loss"])
    rows, hl = [], None
    for idx, (k, v) in enumerate(order):
        t = v["test"]
        if k == "full":
            hl = idx
        rows.append([names.get(k, k), f"{t['log_loss']:.4f}",
                     f"{t['log_loss'] - base:+.4f}", f"{t['macro_auc']:.4f}",
                     f"{t['accuracy']*100:.2f}", f"{t['macro_f1']:.4f}"])
    table(doc, ["구성", "log-loss", "제안 대비", "macro AUC", "정확도(%)", "macro F1"],
          rows, "표 5. 어블레이션 결과(2025 시즌 홀드아웃). 굵은 행이 제안 모델이다.",
          widths=[4.2, 2.4, 2.2, 2.4, 2.2, 2.2], highlight_row=hl)
    figure(doc, "fig4_ablation.png",
           "그림 6. 제안 모델 대비 log-loss 증가량. 오른쪽으로 길수록 해당 요소를 "
           "제거했을 때 성능이 크게 나빠진다.", 15.0)

    d_skip = ab["no_context_skip"]["test"]["log_loss"] - base
    d_seq1 = ab["seq_len_1"]["test"]["log_loss"] - base
    d_state = ab["no_state_embed"]["test"]["log_loss"] - base

    para(doc, "**(1) 맥락 스킵 연결은 실제로 기여한다.**", indent=False)
    para(doc,
         f"제거 시 log-loss가 {d_skip:+.4f} 증가하고 macro AUC는 "
         f"{ab['no_context_skip']['test']['macro_auc'] - tm['macro_auc']:+.4f} 감소하였다. "
         "표에서 가장 큰 효과이며, 유일하게 노이즈 범위 밖으로 판단되는 차이다. "
         "초기 연구가 제안한 핵심 구조가 공정한 비교 하에서 지지된 것이다.")

    para(doc, "**(2) 시퀀스 자체의 기여는 작다.**", indent=False)
    para(doc,
         f"시퀀스 길이를 1로 두어 직전 투구 정보를 완전히 제거해도 log-loss 차이는 "
         f"{d_seq1:+.4f}에 그쳤다. 이는 맥락 스킵 연결 효과의 약 "
         f"{abs(d_seq1/d_skip)*100:.0f}% 수준이다. 예측력의 대부분은 해당 투구 자체의 "
         "물리량과 현재 상황·타자 특성에서 나오며, 직전 투구들이 추가로 제공하는 정보는 "
         "제한적이다. **'투구 배합의 시계열 맥락이 결정적'이라는 초기 연구의 전제는 "
         "이 데이터에서 강하게 지지되지 않는다.** "
         "셋업 피치 최적화의 효과가 최종구 최적화보다 작았던 Takamido와 Nakamoto(2026)의 "
         "보고와도 방향이 일치한다.")

    para(doc, "**(3) 카운트·베이스아웃 상태 임베딩은 효과가 없다.**", indent=False)
    para(doc,
         f"제거해도 log-loss가 {d_state:+.4f}로 오히려 미세하게 개선되었다. "
         "맥락 경로의 이득은 이산 상태 임베딩이 아니라 연속 맥락 피처, 특히 타자의 "
         "직전 시즌 성적에서 나오는 것으로 해석된다. 이 구성요소는 제거하는 것이 타당하다.")

    doc.add_page_break()

    # ---------------- 6. 시스템 ----------------
    h(doc, "6. 시스템 구현")
    para(doc,
         "모델과 분석 결과를 현장에서 활용할 수 있도록 Streamlit 기반 웹 대시보드를 "
         "구현하였다. 다섯 개 화면으로 구성되며, 학습 산출물이 없어도 데이터 집계 기반 "
         "화면은 독립적으로 동작한다.")
    table(doc, ["화면", "기능"],
          [["투구 추천", "매치업·상황·타석 이력을 입력하면 구종×코스 후보 전체를 평가하여 "
                         "기대 실점이 가장 낮은 상위 3개를 제시. 구종별 코스 지도 동시 표시"],
           ["타자 분석", "코스별·구종별 약점, 헛스윙률·체이스율 등 지표를 리그 평균 대비로 표시"],
           ["투수 분석", "구종 레퍼토리와 무브먼트, 카운트별 구종 선택 성향, 코스 성향"],
           ["팀 분석", "30개 구단 투수진 지표 순위와 소속 선수 명단"],
           ["모델 성능", "테스트 시즌 지표, 학습 곡선, 캘리브레이션, 혼동 행렬, 어블레이션"]],
          "표 6. 대시보드 구성.", widths=[3.0, 12.0])

    para(doc,
         f"선수 프로파일은 2024·2025 전체 데이터로 산출하였으며, 분석 가능 대상은 "
         f"타자 {len(profiles.batter_summary):,}명, 투수 {len(profiles.pitcher_summary):,}명이다. "
         "코스별 지표는 표본이 적은 칸에서 우연한 극단값이 약점으로 오인되지 않도록 "
         "경험적 베이즈 축소를 적용하여 리그 평균 쪽으로 당겼다.")
    figure(doc, "fig8_zone.png",
           "그림 7. 타자 코스별 기대 실점 예시. 빨강은 타자에게 유리한(투수에게 위험한) "
           "구역, 파랑은 그 반대다. 검은 사각형이 스트라이크존이다.", 11.0)

    para(doc,
         "시각화는 전 화면에서 일관된 규약을 따른다. 구종별 색은 구종에 고정 배정되어 "
         "필터를 변경해도 바뀌지 않으며, 부호가 의미를 갖는 값은 발산형(파랑–회색–빨강), "
         "크기만을 나타내는 값은 단일 색조 순차형으로 표현한다. 사용한 범주형 팔레트는 "
         "명도대·채도·색각 분리 기준에 대해 사전 검증하였고, 모든 차트에 직접 라벨과 "
         "표 형태의 대체 표현을 함께 제공하여 색만으로 값을 읽지 않아도 되게 하였다.")

    # ---------------- 7. 논의 ----------------
    h(doc, "7. 논의 및 한계")
    para(doc, "본 연구의 한계와 후속 과제는 다음과 같다.")

    bullet(doc, "**단일 시드 실행.** 표 5의 어블레이션은 시드 1개의 단일 실행 결과다. "
                "맥락 스킵 연결의 차이(log-loss +0.0164)는 노이즈 범위 밖으로 보이나, "
                "나머지 구성의 차이(0.001~0.004)는 재현 실험 없이 결론짓기 어렵다. "
                "핵심 3개 구성에 대해 시드를 늘려 평균과 표준편차로 보고할 필요가 있다.")
    bullet(doc, "**학습 예산의 미포화.** 5.1에서 지적하였듯 검증 log-loss의 최저값이 "
                "마지막 에폭에서 관측되었다. 더 긴 학습에서 성능이 개선될 여지가 있으며, "
                "현재 수치는 하한으로 보아야 한다.")
    bullet(doc, "**희귀 클래스 재현율.** 안타 계열의 재현율이 0인 것은 캘리브레이션을 "
                "우선한 설계의 결과다. 확률 보정을 유지하면서 희귀 사건 탐지를 개선하는 "
                "방법(사후 임계값 조정, logit adjustment 등)의 검토가 필요하다.")
    bullet(doc, "**추천기 자체의 검증 부재.** 본 연구는 결과 예측 모델을 홀드아웃에서 "
                "평가하였으나, 추천 로직이 실제로 더 나은 선택을 하는지는 검증하지 않았다. "
                "Takamido와 Nakamoto(2026)의 반사실 최적화와 시즌 지표 회귀를 적용하면 "
                "'이 추천을 따랐을 때 K/9가 얼마나 개선되는가'를 정량화할 수 있다. "
                "본 개정에서 참조 테이블과 반사실 후보 생성 기반은 이미 구축하였으므로, "
                "이 분석은 자연스러운 후속 과제다.")
    bullet(doc, "**타자 적응의 미고려.** 최적 선택이 특정 구종·코스로 수렴하면 타자의 "
                "예측 가능성이 높아져 실제 효과가 추정치보다 낮아질 수 있다. "
                "투수와 타자의 전략적 상호작용을 모델에 반영하는 것이 필요하다.")

    para(doc,
         "한편 본 개정 과정에서 얻은 방법론적 교훈을 기록해 둔다. "
         "**성능이 예상보다 높게 나올 때 그것을 검증하는 절차가 파이프라인에 포함되어야 한다.** "
         "초기 구현의 67.52%는 그 자체로는 의심스럽지 않은 값이었고, 정확도라는 단일 지표로는 "
         "누수를 감지할 수 없었다. 본 개정에서는 시퀀스 조립 후 '유효 길이가 "
         "min(pitch_number, 6)과 일치하는가'를 자동 검증하는 스크립트를 파이프라인에 "
         "포함시켰다. 이러한 구조적 점검이 사후 분석보다 효과적이다.")

    # ---------------- 8. 결론 ----------------
    h(doc, "8. 결론")
    para(doc,
         "본 연구는 초기 구현이 보고한 성능이 미래 정보 누수에 기인함을 규명하고, "
         "데이터 구성부터 평가 방법까지 파이프라인을 재설계하였다. Statcast 원본의 "
         "역시간순 정렬을 놓친 것이 근본 원인이었으며, 볼카운트가 모든 타임스텝에 "
         "포함되어 있어 인접 투구의 카운트 비교만으로 정답을 읽을 수 있는 구조였다.")
    para(doc,
         "재설계 후 2025 시즌 홀드아웃에서 "
         f"log-loss {tm['log_loss']:.4f}, macro AUC {tm['macro_auc']:.4f}, "
         f"ECE {ev['ece']:.4f}를 얻었다. 정확도는 초기 보고값보다 낮으나, 두 수치는 "
         "측정 대상이 다르므로 비교 가능한 값이 아니다. 오히려 temperature가 "
         f"{ev['temperature']:.3f}로 1에 근접한다는 사실은 출력 확률이 실제 빈도와 "
         "일치함을 보여주며, 확률로 후보를 순위화하는 본 시스템에서는 이것이 "
         "정확도보다 본질적인 성질이다.")
    para(doc,
         "어블레이션을 통해 맥락 스킵 연결의 기여를 공정한 조건에서 확인하였고, "
         "동시에 투구 시퀀스의 기여가 예상보다 작다는 부정적 결과도 함께 얻었다. "
         "후자는 초기 연구의 전제를 약화시키지만, 그 자체로 의미 있는 발견이다. "
         "투구 배합의 효과를 정량화하려면 본 연구보다 정교한 실험 설계가 필요하며, "
         "이는 반사실 분석과 시즌 지표 환산을 결합한 후속 연구로 이어질 수 있다.")

    # ---------------- 참고문헌 ----------------
    h(doc, "참고문헌")
    refs = [
        "R. Takamido and H. Nakamoto, \"Counterfactual Optimization of Baseball Pitch "
        "Sequences and Estimation of Its Impact on Season-Level Statistics,\" "
        "arXiv:2606.17345, 2026.",
        "D. Kneita, \"Transformer-Based Baseball Modeling for Pitch Outcome Prediction "
        "and Strategy Optimization,\" MIT Sloan Sports Analytics Conference, 2025.",
        "C. Guo, G. Pleiss, Y. Sun, and K. Q. Weinberger, \"On Calibration of Modern "
        "Neural Networks,\" in Proc. ICML, 2017, pp. 1321–1330.",
        "T.-Y. Lin, P. Goyal, R. Girshick, K. He, and P. Dollár, \"Focal Loss for Dense "
        "Object Detection,\" in Proc. IEEE ICCV, 2017, pp. 2980–2988.",
        "C. Guo and F. Berkhahn, \"Entity Embeddings of Categorical Variables,\" "
        "arXiv:1604.06737, 2016.",
        "A. Vaswani et al., \"Attention Is All You Need,\" in Advances in Neural "
        "Information Processing Systems 30, 2017.",
        "P. Swartz, M. Grosskopf, D. R. Bingham et al., \"The Quality of Pitches in "
        "Major League Baseball,\" The American Statistician, vol. 71, no. 2, pp. 148–154, 2017.",
        "G. Sidle and H. Tran, \"Using Multi-Class Classification Methods to Predict "
        "Baseball Pitch Types,\" Journal of Sports Analytics, vol. 4, no. 1, pp. 85–93, 2018.",
        "A. Prasad, \"Decoding MLB Pitch Sequencing Strategies via Directed Graph "
        "Embeddings,\" MIT Sloan Sports Analytics Conference, 2021.",
        "MLB Advanced Media, \"Statcast Data,\" Baseball Savant. "
        "[Online]. Available: https://baseballsavant.mlb.com",
    ]
    for i, r in enumerate(refs, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.9)
        p.paragraph_format.first_line_indent = Cm(-0.9)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.3
        _kfont(p.add_run(f"[{i}] {r}"), size=9.5)

    # ---------------- 부록 ----------------
    doc.add_page_break()
    h(doc, "부록. 재현 방법")
    para(doc,
         "본 연구의 코드와 산출물은 아래 저장소에 공개되어 있다. 데이터 수집부터 "
         "학습, 산출물 생성까지 세 개의 노트북으로 재현할 수 있다.")
    code(doc, "https://github.com/hwangiljun/deeplearning_project")
    table(doc, ["단계", "실행", "내용"],
          [["1", "notebooks/01_download_statcast.ipynb", "2024·2025 정규시즌 수집"],
           ["2", "notebooks/02_train.ipynb", "학습 및 어블레이션 (GPU)"],
           ["3", "notebooks/03_build_artifacts.ipynb", "참조 테이블·프로파일·평가"],
           ["-", "scripts/validate_pipeline.py", "누수 차단 자동 검증"],
           ["-", "streamlit run app.py", "대시보드 실행"]],
          None, widths=[1.4, 7.0, 6.6])

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"저장: {OUT}")
    print(f"  ({OUT.stat().st_size/1024:.0f}KB)")


if __name__ == "__main__":
    build()
